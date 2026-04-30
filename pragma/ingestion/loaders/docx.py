from pathlib import Path
from typing import Any, List

from pragma.ingestion.loader import DocumentSegment


def load_docx_file(path: Path) -> List[DocumentSegment]:
    """Load DOCX file."""
    try:
        import docx
    except ImportError:
        return []

    doc = docx.Document(str(path))
    paragraphs = []
    heading_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            heading_parts.append(text)
        else:
            paragraphs.append(text)

    tables_content = _extract_tables(doc)

    full_content = "\n\n".join(paragraphs)
    if tables_content:
        full_content += "\n\n" + tables_content

    return [
        DocumentSegment(
            content=full_content,
            source=str(path),
            doc_type="docx",
            metadata={
                "filename": path.name,
                "headings": heading_parts,
                "paragraph_count": len(paragraphs),
            },
        )
    ]


def _extract_tables(doc: Any) -> str:
    """Extract tables from DOCX."""
    if not hasattr(doc, "tables") or not doc.tables:
        return ""

    parts = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        parts.append("\n".join(rows))

    return "\n\n".join(parts)
